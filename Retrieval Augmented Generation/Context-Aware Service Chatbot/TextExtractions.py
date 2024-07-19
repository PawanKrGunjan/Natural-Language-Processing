"""
pip install language-tool-python
pip install textblob
pip install langchain_community
pip install pytesseract
pip install pyspellchecker
pip install spacy
"""

from tqdm import tqdm
import os
from bs4 import BeautifulSoup
import re
import language_tool_python
#from textblob import TextBlob
from docx import Document
import cv2
from PIL import Image
import pytesseract
import spacy
from spellchecker import SpellChecker
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
import langchain_core
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_groq import ChatGroq
from langchain_cohere import CohereEmbeddings
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
import json
from typing import Literal, Dict


class SimpleDocument:
    def __init__(self, page_content):
        """Initialize SimpleDocument with page content."""
        self.page_content = page_content
        self.metadata = {}

class RetrievalAugmentedGeneration:
    def __init__(self,secrets_file):
        """
        Initialize the TextExtractions
        """
        self.keys = self._load_secret_keys(secrets_file)
        self._setup_environment()
        self.documents = []
        self.nlp = spacy.load("en_core_web_sm")
        self.spellChecker = SpellChecker()
        self.tool = language_tool_python.LanguageTool('en-US')
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=250)
        self.embedding = CohereEmbeddings()
        self.llm = ChatGroq(model="llama3-8b-8192")

    def _load_secret_keys(self, file_path):
        """Load secret keys from a JSON file."""
        with open(file_path, 'r') as file:
            secret_keys = json.load(file)
        return secret_keys

    def _setup_environment(self):
        """Setup environment variables for LangChain, Cohere, and Groq."""
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_API_KEY"] = self.keys['Langchain Smith']
        os.environ["COHERE_API_KEY"] = self.keys['cohere new']
        os.environ['OPENAI_API_KEY']=self.keys['Open ai']
        os.environ["GROQ_API_KEY"] = self.keys['Groq']

    def clean_text(self, text):
        """Clean text by removing HTML tags, excessive whitespace, and other unnecessary characters."""
    
        # Remove HTML tags
        soup = BeautifulSoup(text, "html.parser")
        text = soup.get_text()
    
        # Remove specific unwanted characters
        text = re.sub(r'\x83', '', text)  # Removing specific unwanted characters
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Removing non-ASCII characters
        text = re.sub(r'(?<!\S)-{2,}(?!\S)', ' ', text)  # Removing multiple hyphens not within words
    
        # Separate headings (considering headings are likely to be in uppercase or have newlines around them)
        text = re.sub(r'(_\n\s*[A-Z][A-Z0-9\s\-]*\n)', r'\n\1\n', text)
        
        # Replace multiple spaces with a single space, but keep newlines
        text = re.sub(r'\s+', ' ', text) # Replace multiple spaces with a single space
        text = re.sub(r'\n+', '\n', text)  # Replace multiple newlines with a single newline

        # Remove punctuation that doesn't serve a purpose
        text = re.sub(r'[^\w\s\n]', '', text)  # Remove all punctuation except for newlines and words

        # Remove underscores
        text = re.sub(r'_', ' ', text)  # Replace underscores with a space
    
        # Remove excessive whitespace
        text = text.strip()
    
        return text

    def spell_check(self, text):
        # Process the text with SpaCy
        doc = self.nlp(text)
    
        # Spell checking
        corrected_text = []
        for token in doc:
            if token.is_alpha:  # Check if the token is a word
                # Get candidates from spell checker
                candidates = self.spellChecker.candidates(token.text)
                # Use the first candidate or the original word if no candidates are available
                corrected_word = next(iter(candidates), token.text) if candidates else token.text
                corrected_text.append(corrected_word)
            else:
                corrected_text.append(token.text)
    
        return ' '.join(corrected_text)

    def grammar_style_check(self, text):
        """Perform grammar and style check"""
        matches = self.tool.check(text)
        return language_tool_python.utils.correct(text, matches)

    def preprocess_image_and_extract_text(self, image_path):
        """
        Preprocess an image and extract text using OCR.
    
        :param image_path: Path to the input image file.
        :param preprocessed_image_path: Path to save the preprocessed image.
        :return: Extracted text from the image.
        """
        try:
            # Load the image using OpenCV
            image = cv2.imread(image_path)
        
            # Convert the image to grayscale
            gray_image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
            # Apply thresholding
            _, binary_image = cv2.threshold(gray_image, 128, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
        
            preprocessed_image_path='temp.png'
        
            # Save the preprocessed image
            cv2.imwrite(preprocessed_image_path, binary_image)
        
            # Open the preprocessed image using PIL
            preprocessed_image = Image.open(preprocessed_image_path)
        
            # Use pytesseract to do OCR on the preprocessed image
            text = pytesseract.image_to_string(preprocessed_image)
    
            processed_text = self.clean_text(text)
            processed_text = self.spell_check(processed_text)
            processed_text = self.grammar_style_check(processed_text)
            simple_doc = SimpleDocument(processed_text)
            chunks = self.text_splitter.split_documents([simple_doc])
            for idx, chunk in enumerate(chunks):
                chunk.metadata['source'] = image_path.split('/')[-1]
                chunk.metadata['page'] =  0
                self.documents.append(chunk)
        except Exception as e:
            print(f"Failed to extract text from {image_path.split('/')[-1]}: {e}")
        return f"Text extraction from {image_path.split('/')[-1]} completed."

    def load_text_file(self, file_path):
        """Load and store content from a text file."""
        with open(file_path, 'r') as file:
            text = file.read()
            processed_text = self.clean_text(text)
            processed_text = self.spell_check(processed_text)
            processed_text = self.grammar_style_check(processed_text)
            simple_doc = SimpleDocument(processed_text)
            chunks = self.text_splitter.split_documents([simple_doc])
            for idx, chunk in enumerate(chunks):
                chunk.metadata['source'] = file_path.split('/')[-1]
                chunk.metadata['page'] =  0
                self.documents.append(chunk)
        return f"Text extraction from {file_path.split('/')[-1]} completed."

    def load_pdf_file(self, file_path):
        """
        Load and store content from a PDF file.
        """
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        doc_content = '\n'.join([page.page_content for page in tqdm(docs)])
        i = 1
        for page in tqdm(docs, desc="Reading PDF file"):
            try:
                text = page.page_content
                processed_text = self.clean_text(text)
                #processed_text = self.spell_check(processed_text)
                #processed_text = self.grammar_style_check(processed_text)
                simple_doc = SimpleDocument(processed_text)
                chunks = self.text_splitter.split_documents([simple_doc])
                for idx, chunk in enumerate(chunks):
                    chunk.metadata['source'] = f"{file_path.split('/')[-1]} Page No. {i}"
                    chunk.metadata['page'] =  i
                    self.documents.append(chunk)
                i += 1
            except Exception as e:
                print(f"Failed to extract text from {file_path.split('/')[-1]}: {e}")
        return f"Text extraction from {file_path.split('/')[-1]} completed."

    def load_docx_file(self, file_path):
        """
        Load and store content from a DOCX file.
        """
        try:
            doc = Document(file_path)
            full_text = []
            title = doc.core_properties.title
            if title:
                full_text.append(title)
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        correctly_spelled = self.spell_check(cell.text)
                        corrected_texts = self.grammar_style_check(correctly_spelled)
                        full_text.append(corrected_texts)
            document_text = '\n'.join(full_text)
            simple_doc = SimpleDocument(document_text)
            chunks = self.text_splitter.split_documents([simple_doc])
            for idx, chunk in enumerate(chunks):
                chunk.metadata['source'] = file_path.split('/')[-1]
                chunk.metadata['page'] =  0
                self.documents.append(chunk)
                self.documents[file_path.split('/')[-1]] = document_text
        except Exception as e:
            print(f"Failed to extract text from {file_path.split('/')[-1]}: {e}")

    def load_documents(self, file_paths):
        """
        Load documents from various file types.
        """
        for file_path in tqdm(file_paths, desc="Loading documents"):
            extension = os.path.splitext(file_path)[1].lower()
            if extension == '.txt':
                self.load_text_file(file_path)
            elif extension == '.pdf':
                self.load_pdf_file(file_path)
            elif extension == '.docx':
                self.load_docx_file(file_path)
            elif extension in ['.jpg', '.png', '.jpeg']:
                self.preprocess_image_and_extract_text(file_path)
            else:
                print(f"Unsupported file type: {extension}")

    
    def text_vectorizations(self):
        """Generate text embeddings and create a retriever for document retrieval."""
        vectorstore = Chroma.from_documents(documents=self.documents, embedding=self.embedding)
        retriever = vectorstore.as_retriever()
        return retriever
    
    def create_prompt(self, system_template):
        """Create a prompt template for the language model."""
        prompt = PromptTemplate(
            input_variables=["context", "question"],
            template_format='f-string',
            template=system_template,
            output_parser=None,
            validate_template=True
        )
        return prompt
    
    def rag_chain(self,system_template):
        """Create a Retrieval-Augmented Generation (RAG) chain for question answering."""
        rag_chain=RetrievalQA.from_chain_type(self.llm, 
                                              retriever=self.text_vectorizations(), 
                                              chain_type_kwargs={"prompt": self.create_prompt(system_template)})
        return rag_chain
    
    def question_answering(self, question,system_template):
        """Answer a question using the RAG chain."""
        rag_chain = self.rag_chain(system_template)
        response = rag_chain.invoke(question)
        return response["result"]

chatbot = RetrievalAugmentedGeneration('secrets.json')

massages=chatbot.load_documents(['Help files India DECU machines.pdf'])
prompt_template = """
<|start_header_id|>user<|end_header_id|>
You are an assistant to provide the help for the query based on the given document in as below format.

First Identify the main issues or problem by the given DTC code or exact issues.
DTC code is the unique identity number designated for the particular type of issues.
Then provide the result as per the below format clearly.

DTC code: The issue numbere designated for the particular problem with short one line descriptions.
Possible Cause: Discuss the possible cause pointwise.
Diagnosis: Discuss how to perform diagnosis stepwise for the given problem or issue number.
Corrective Action: Discuss how to take preventive measure or how to deal with that particular isses stepwise to correct it.

**Note**: If no relevant context or answer is found in the provided documents, please specify that no result has been found.
--------------------------------
Question: {question}
Context: {context}
"""

# Example query
question = "C131A-64: Lever lock switches Plausibility Error"

# Get the response
response = chatbot.question_answering(question,prompt_template)
print(response)

# Example query
question = "C131A-64: Lever lock Error"

# Get the response
response = chatbot.question_answering(question,prompt_template)
print(response)

# Example query
question = "C131A-64 Error"

# Get the response
response = chatbot.question_answering(question,prompt_template)
print(response)

# Example query
question = "C131A-64 Error"

# Get the response
response = chatbot.question_answering(question,prompt_template)
print(response)