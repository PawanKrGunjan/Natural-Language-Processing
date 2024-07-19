"""
json==2.0.9
re==2.2.1
beautifulsoup4==4.12.3
cryptography==42.0.8
langchain-community==0.2.6
python-docx==1.1.2
langchain==0.2.6
langchain-core==0.2.11
langchain-groq
"""

import os
import json
import re
from bs4 import BeautifulSoup
from langchain_community.document_loaders import PyPDFLoader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
#from langchain.embeddings import CohereEmbeddings
from langchain_cohere import CohereEmbeddings
#from langchain_openai import OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_groq import ChatGroq
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain

class SimpleDocument:
    def __init__(self, page_content):
        """Initialize SimpleDocument with page content."""
        self.page_content = page_content
        self.metadata = {}

class RetrievalAugmentedGeneration:
    def __init__(self, secrets_file, prompt_template):
        """Initialize the RetrievalAugmentedGeneration with secret keys and prompt template."""
        self.keys = self._load_secret_keys(secrets_file)
        self._setup_environment()
        self.documents = {}
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=250)
        self.embedding = CohereEmbeddings()
        self.system_prompt = prompt_template
        self.llm = ChatGroq(model="llama3-8b-8192")

    def _load_secret_keys(self, file_path):
        """Load secret keys from a JSON file."""
        with open(file_path, 'r') as file:
            secret_keys = json.load(file)
        return secret_keys

    def _setup_environment(self):
        """Setup environment variables for LangChain, Cohere, and Groq."""
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_API_KEY"] = self.keys['Langchain Smith']
        os.environ["COHERE_API_KEY"] = self.keys['cohere new']
        os.environ['OPENAI_API_KEY']=self.keys['Open ai']
        os.environ["GROQ_API_KEY"] = self.keys['Groq']

    def load_text_file(self, file_path):
        """Load and store content from a text file."""
        with open(file_path, 'r') as file:
            text = file.read()
            self.documents[file_path.split('/')[-1]] = text
        return f"Text extraction from {file_path.split('/')[-1]} completed."

    def load_pdf_file(self, file_path):
        """Load and store content from a PDF file."""
        try:
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            doc_content = '\n'.join([page.page_content for page in docs])
            self.documents[file_path.split('/')[-1]] = doc_content
            return f"Text extraction from {file_path.split('/')[-1]} completed."
        except Exception as e:
            return f"Failed to extract text from {file_path.split('/')[-1]}: {e}"

    def load_docx_file(self, file_path):
        """Load and store content from a DOCX file."""
        try:
            doc = Document(file_path)
            full_text = []
            title = doc.core_properties.title
            if title:
                full_text.append(title)
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            document_text = '\n'.join(full_text)
            self.documents[file_path.split('/')[-1]] = document_text
            return f"Text extraction from {file_path.split('/')[-1]} completed."
        except Exception as e:
            return f"Failed to extract text from {file_path.split('/')[-1]}: {e}"

    def load_documents(self, file_paths):
        """Load documents from various file types."""
        massages =[]
        for file_path in file_paths:
            extension = os.path.splitext(file_path)[1].lower()
            if extension == '.txt':
                massages.append(self.load_text_file(file_path))
            elif extension == '.pdf':
                massages.append(self.load_pdf_file(file_path))
            elif extension == '.docx':
                massages.append(self.load_docx_file(file_path))
            else:
                massages.append(f"Unsupported file type: {extension}")
        return massages

    def get_all_documents(self):
        """Retrieve all loaded documents."""
        return self.documents

    def clean_text(self, text):
        """Clean text by removing HTML tags, excessive whitespace, and other unnecessary characters."""
        soup = BeautifulSoup(text, "html.parser")
        text = soup.get_text()
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\*\*|###|\[[^]]*\]', ' ', text)
        text = text.strip()
        text = re.sub(r'\n+', '\n', text)
        return text

    def preprocessing(self):
        """Preprocess documents by cleaning and splitting them into chunks."""
        preprocessed_chunks = []
        for doc_name, doc_content in self.get_all_documents().items():
            cleaned_content = self.clean_text(doc_content)
            simple_doc = SimpleDocument(cleaned_content)
            chunks = self.text_splitter.split_documents([simple_doc])
            for idx, chunk in enumerate(chunks):
                chunk.metadata['source'] = doc_name
                chunk.metadata['page'] = idx
                preprocessed_chunks.append(chunk)
        return preprocessed_chunks

    def text_vectorizations(self):
        """Generate text embeddings and create a retriever for document retrieval."""
        vectorstore = Chroma.from_documents(documents=self.preprocessing(), embedding=self.embedding)
        retriever = vectorstore.as_retriever()
        return retriever

    def _create_prompt(self):
        """Create a prompt template for the language model."""
        prompt = ChatPromptTemplate.from_messages(
            [
                ("system", self.system_prompt),
                ("human", "{input}"),
            ])
        return prompt

    def _create_rag_chain(self):
        """Create a Retrieval-Augmented Generation (RAG) chain for question answering."""
        question_answer_chain = create_stuff_documents_chain(self.llm, self._create_prompt())
        retriever = self.text_vectorizations()
        rag_chain = create_retrieval_chain(retriever, question_answer_chain)
        return rag_chain

    def question_answering(self, question):
        """Answer a question using the RAG chain."""
        rag_chain = self._create_rag_chain()
        response = rag_chain.invoke({"input": question})
        return response
