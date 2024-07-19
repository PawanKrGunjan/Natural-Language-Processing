import os
import json
import re
from bs4 import BeautifulSoup
from langchain_community.document_loaders import PyPDFLoader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_cohere import CohereEmbeddings
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_groq import ChatGroq
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from tqdm import tqdm

class SimpleDocument:
    def __init__(self, page_content):
        """Initialize SimpleDocument with page content."""
        self.page_content = page_content
        self.metadata = {}

class RetrievalAugmentedGeneration:
    def __init__(self, secrets_file, prompt_template):
        """Initialize the RetrievalAugmentedGeneration with secret keys and prompt template."""
        self.keys = self._load_secret_keys(secrets_file)
        self._setup_environment()
        self.documents = {}
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=250)
        self.embedding = CohereEmbeddings()
        self.system_prompt = prompt_template
        self.llm = ChatGroq(model="llama3-8b-8192")
        self.history = []  # Initialize an empty history list

    def _load_secret_keys(self, file_path):
        """Load secret keys from a JSON file."""
        with open(file_path, 'r') as file:
            secret_keys = json.load(file)
        return secret_keys

    def _setup_environment(self):
        """Setup environment variables for LangChain, Cohere, and Groq."""
        os.environ["LANGCHAIN_TRACING_V2"] = "true"
        os.environ["LANGCHAIN_API_KEY"] = self.keys['Langchain Smith']
        os.environ["COHERE_API_KEY"] = self.keys['cohere new']
        os.environ['OPENAI_API_KEY'] = self.keys['Open ai']
        os.environ["GROQ_API_KEY"] = self.keys['Groq']

    def clean_text(self, text):
        """Clean text by removing HTML tags, excessive whitespace, and other unnecessary characters."""
        soup = BeautifulSoup(text, "html.parser")
        text = soup.get_text()
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\*\*|###|\[[^]]*\]', ' ', text)
        text = text.strip()
        text = re.sub(r'\n', ' ', text)
        text = re.sub(r'\n+', ' ', text)
        return text

    def load_text_file(self, file_path):
        """Load and store content from a text file."""
        with open(file_path, 'r') as file:
            text = file.read()
            cleaned_text = self.clean_text(text)
            self.documents[file_path.split('/')[-1]] = cleaned_text
        return f"Text extraction from {file_path.split('/')[-1]} completed."

    def load_pdf_file(self, file_path):
        """Load and store content from a PDF file."""
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        i = 1
        for page in tqdm(docs):
            try:
                text = page.page_content
                cleaned_text = self.clean_text(text)
                self.documents[f"{file_path.split('/')[-1]} 'Page No. {i}"] = cleaned_text
                #print(f"Text extraction from {file_path.split('/')[-1]} Page No. {i} completed.")
                i += 1
            except Exception as e:
                print(f"Failed to extract text from {file_path.split('/')[-1]} Page No. {i}: {e}")

    def load_docx_file(self, file_path):
        """Load and store content from a DOCX file."""
        try:
            doc = Document(file_path)
            full_text = []
            title = doc.core_properties.title
            if title:
                full_text.append(title)
            for paragraph in doc.paragraphs:
                cleaned_text = self.clean_text(paragraph.text)
                full_text.append(cleaned_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cleaned_text = self.clean_text(cell.text)
                        full_text.append(cleaned_text)
            document_text = '\n'.join(full_text)
            self.documents[file_path.split('/')[-1]] = document_text
            return f"Text extraction from {file_path.split('/')[-1]} completed."
        except Exception as e:
            return f"Failed to extract text from {file_path.split('/')[-1]}: {e}"

    def load_documents(self, file_paths):
        """Load documents from various file types."""
        messages = []
        for file_path in file_paths:
            extension = os.path.splitext(file_path)[1].lower()
            if extension == '.txt':
                messages.append(self.load_text_file(file_path))
            elif extension == '.pdf':
                messages.append(self.load_pdf_file(file_path))
            elif extension == '.docx':
                messages.append(self.load_docx_file(file_path))
            else:
                messages.append(f"Unsupported file type: {extension}")
        return messages

    def get_all_documents(self):
        """Retrieve all loaded documents."""
        return self.documents

    def preprocessing(self):
        """Preprocess documents by cleaning and splitting them into chunks."""
        preprocessed_chunks = []
        for doc_name, doc_content in self.get_all_documents().items():
            simple_doc = SimpleDocument(doc_content)
            chunks = self.text_splitter.split_documents([simple_doc])
            for idx, chunk in enumerate(chunks):
                chunk.metadata['source'] = doc_name
                chunk.metadata['page'] = int(doc_name.split()[-1]) if doc_name.split()[-1].isdigit() else None
                preprocessed_chunks.append(chunk)
        return preprocessed_chunks

    def text_vectorizations(self):
        """Generate text embeddings and create a retriever for document retrieval."""
        vectorstore = Chroma.from_documents(documents=self.preprocessing(), embedding=self.embedding)
        retriever = vectorstore.as_retriever()
        return retriever

    def _create_prompt(self):
        """Create a prompt template for the language model."""
        prompt = ChatPromptTemplate.from_messages(
            [
                ("system", self.system_prompt),
                MessagesPlaceholder(variable_name="history"),
                ("human", "{input}"),
            ])
        return prompt

    def _create_rag_chain(self):
        """Create a Retrieval-Augmented Generation (RAG) chain for question answering."""
        question_answer_chain = create_stuff_documents_chain(self.llm, self._create_prompt())
        retriever = self.text_vectorizations()
        rag_chain = create_retrieval_chain(retriever, question_answer_chain)
        return rag_chain

    def question_answering(self, question):
        """Answer a question using the RAG chain and update the history."""
        rag_chain = self._create_rag_chain()
        # Update the history with the current question and get the response
        response = rag_chain.invoke({"input": question, "history": self.history})
        print(response)
        # Update the history with the current question and response
        self.history.append({"role": "user", "content": question})
        self.history.append({"role": "assistant", "content": response['answer']})
        return response['answer']

# Example usage
prompt_template = """
You are an assistant for finding the exact or similar descriptions to questions using the provided context. 
If the question is not related to the provided context, please specify. 
Please provide further details or related references for your question. 
Keep the answer concise.

{context}
"""

bot = RetrievalAugmentedGeneration('secrets.json', prompt_template)

bot.load_documents(['Help files India DECU machines.pdf'])
response = bot.question_answering('Hi, My name is PawanKumarGunjan? I want to know about DECU machines?')
print(response)

response = bot.question_answering('What is my name?')
print(response)
